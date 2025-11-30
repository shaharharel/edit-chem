"""
Report generator for RNA MPRA experiments.

Generates comprehensive analysis reports similar to small_molecules/report_generator.py
but adapted for RNA sequence edit prediction.

Features:
- Configuration summary
- Method comparison plots (MAE, RMSE, R², Pearson r, Direction Accuracy)
- Scatter plots (predicted vs actual)
- Delta distribution analysis
- Edit position analysis
- Mutation type breakdown
- Cell type comparison (if multiple)
- Training progress visualization
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

import json
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from scipy import stats

# Try to import docx for Word reports
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("python-docx not installed. Word report generation disabled.")


def set_plot_style():
    """Set consistent plot style for all figures."""
    plt.style.use('seaborn-v0_8-whitegrid')
    plt.rcParams.update({
        'font.size': 11,
        'axes.titlesize': 14,
        'axes.labelsize': 12,
        'xtick.labelsize': 10,
        'ytick.labelsize': 10,
        'legend.fontsize': 10,
        'figure.titlesize': 16,
        'figure.dpi': 150
    })


def create_metrics_comparison_plot(
    results: Dict[str, Dict],
    metrics: List[str],
    output_path: str,
    title: str = "Method Comparison"
) -> str:
    """
    Create bar plot comparing metrics across methods.

    Args:
        results: Dict of {method_name: {metric: value}}
        metrics: List of metrics to plot
        output_path: Path to save the plot
        title: Plot title

    Returns:
        Path to saved plot
    """
    set_plot_style()

    methods = list(results.keys())
    n_metrics = len(metrics)

    fig, axes = plt.subplots(1, n_metrics, figsize=(4 * n_metrics, 5))
    if n_metrics == 1:
        axes = [axes]

    # Color palette
    colors = plt.cm.Set2(np.linspace(0, 1, len(methods)))

    for idx, metric in enumerate(metrics):
        ax = axes[idx]
        values = [results[m].get(metric, 0) for m in methods]

        bars = ax.bar(range(len(methods)), values, color=colors)
        ax.set_xticks(range(len(methods)))
        ax.set_xticklabels(methods, rotation=45, ha='right')
        ax.set_ylabel(metric.upper().replace('_', ' '))
        ax.set_title(metric.upper().replace('_', ' '))

        # Add value labels on bars
        for bar, val in zip(bars, values):
            height = bar.get_height()
            ax.annotate(f'{val:.3f}',
                       xy=(bar.get_x() + bar.get_width() / 2, height),
                       xytext=(0, 3),
                       textcoords="offset points",
                       ha='center', va='bottom', fontsize=9)

        # Add baseline reference line for direction accuracy
        if metric == 'direction_accuracy':
            ax.axhline(y=0.5, color='red', linestyle='--', alpha=0.7, label='Random (50%)')
            ax.legend()

        # For R², add zero line
        if metric == 'r2':
            ax.axhline(y=0, color='gray', linestyle='--', alpha=0.5)

    plt.suptitle(title, fontsize=14, fontweight='bold')
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

    return output_path


def create_scatter_plot(
    y_true: np.ndarray,
    y_pred: np.ndarray,
    method_name: str,
    output_path: str,
    include_metrics: bool = True
) -> str:
    """
    Create scatter plot of predicted vs actual values.

    Args:
        y_true: True delta values
        y_pred: Predicted delta values
        method_name: Name of the method
        output_path: Path to save the plot
        include_metrics: Whether to include metrics in the plot

    Returns:
        Path to saved plot
    """
    set_plot_style()

    fig, ax = plt.subplots(figsize=(8, 8))

    # Scatter plot with transparency
    ax.scatter(y_true, y_pred, alpha=0.5, s=30, c='steelblue', edgecolors='none')

    # Perfect prediction line
    min_val = min(y_true.min(), y_pred.min())
    max_val = max(y_true.max(), y_pred.max())
    ax.plot([min_val, max_val], [min_val, max_val], 'r--', lw=2, label='Perfect Prediction')

    # Zero lines
    ax.axhline(y=0, color='gray', linestyle='-', alpha=0.3)
    ax.axvline(x=0, color='gray', linestyle='-', alpha=0.3)

    # Fit line
    if len(y_true) > 2:
        z = np.polyfit(y_true, y_pred, 1)
        p = np.poly1d(z)
        x_line = np.linspace(min_val, max_val, 100)
        ax.plot(x_line, p(x_line), 'g-', alpha=0.7, lw=1.5, label=f'Fit: y={z[0]:.2f}x+{z[1]:.2f}')

    ax.set_xlabel('True Δ (MRL)', fontsize=12)
    ax.set_ylabel('Predicted Δ (MRL)', fontsize=12)
    ax.set_title(f'{method_name}\nPredicted vs Actual', fontsize=14, fontweight='bold')
    ax.legend(loc='upper left')

    # Add metrics text box
    if include_metrics:
        mae = np.mean(np.abs(y_true - y_pred))
        rmse = np.sqrt(np.mean((y_true - y_pred) ** 2))
        r2 = 1 - np.sum((y_true - y_pred) ** 2) / np.sum((y_true - np.mean(y_true)) ** 2)
        pearson_r = stats.pearsonr(y_true, y_pred)[0]

        textstr = f'MAE: {mae:.4f}\nRMSE: {rmse:.4f}\nR²: {r2:.4f}\nPearson r: {pearson_r:.4f}'
        props = dict(boxstyle='round', facecolor='wheat', alpha=0.8)
        ax.text(0.95, 0.05, textstr, transform=ax.transAxes, fontsize=10,
                verticalalignment='bottom', horizontalalignment='right', bbox=props)

    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

    return output_path


def create_delta_distribution_plot(
    data: pd.DataFrame,
    predictions: Optional[Dict[str, np.ndarray]] = None,
    output_path: str = None
) -> str:
    """
    Create distribution plot of delta values.

    Args:
        data: DataFrame with 'delta' column
        predictions: Optional dict of {method_name: predicted_deltas}
        output_path: Path to save the plot

    Returns:
        Path to saved plot
    """
    set_plot_style()

    n_plots = 1 + (len(predictions) if predictions else 0)
    fig, axes = plt.subplots(1, min(n_plots, 3), figsize=(5 * min(n_plots, 3), 5))
    if n_plots == 1:
        axes = [axes]

    # True delta distribution
    ax = axes[0]
    ax.hist(data['delta'], bins=50, color='steelblue', alpha=0.7, edgecolor='black')
    ax.axvline(x=0, color='red', linestyle='--', alpha=0.7)
    ax.axvline(x=data['delta'].mean(), color='green', linestyle='-', alpha=0.7,
               label=f'Mean: {data["delta"].mean():.3f}')
    ax.set_xlabel('Δ MRL')
    ax.set_ylabel('Count')
    ax.set_title('True Delta Distribution')
    ax.legend()

    # Predicted distributions (if provided)
    if predictions:
        for i, (method, preds) in enumerate(list(predictions.items())[:2]):
            ax = axes[i + 1]
            ax.hist(preds, bins=50, color='coral', alpha=0.7, edgecolor='black')
            ax.axvline(x=0, color='red', linestyle='--', alpha=0.7)
            ax.axvline(x=np.mean(preds), color='green', linestyle='-', alpha=0.7,
                      label=f'Mean: {np.mean(preds):.3f}')
            ax.set_xlabel('Δ MRL (Predicted)')
            ax.set_ylabel('Count')
            ax.set_title(f'{method}\nPredicted Delta Distribution')
            ax.legend()

    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

    return output_path


def create_mutation_type_analysis(
    data: pd.DataFrame,
    results: Optional[Dict[str, Dict]] = None,
    output_path: str = None
) -> Tuple[str, pd.DataFrame]:
    """
    Analyze performance by mutation type (SNV).

    Args:
        data: DataFrame with edit_from and edit_to columns
        results: Optional results dict with predictions
        output_path: Path to save the plot

    Returns:
        Tuple of (plot_path, analysis_df)
    """
    set_plot_style()

    # Create mutation type column
    data = data.copy()
    data['mutation_type'] = data['edit_from'] + '→' + data['edit_to']

    # Group by mutation type
    mutation_stats = data.groupby('mutation_type').agg({
        'delta': ['mean', 'std', 'count']
    }).round(4)
    mutation_stats.columns = ['mean_delta', 'std_delta', 'count']
    mutation_stats = mutation_stats.reset_index()
    mutation_stats = mutation_stats.sort_values('mean_delta', ascending=False)

    # Create plot
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))

    # Mean delta by mutation type
    ax = axes[0]
    colors = ['coral' if x > 0 else 'steelblue' for x in mutation_stats['mean_delta']]
    bars = ax.bar(mutation_stats['mutation_type'], mutation_stats['mean_delta'], color=colors)
    ax.axhline(y=0, color='black', linestyle='-', alpha=0.5)
    ax.set_xlabel('Mutation Type')
    ax.set_ylabel('Mean Δ MRL')
    ax.set_title('Mean Effect by Mutation Type', fontweight='bold')
    ax.tick_params(axis='x', rotation=45)

    # Add error bars
    ax.errorbar(mutation_stats['mutation_type'], mutation_stats['mean_delta'],
                yerr=mutation_stats['std_delta'], fmt='none', color='black', alpha=0.5)

    # Count by mutation type
    ax = axes[1]
    ax.bar(mutation_stats['mutation_type'], mutation_stats['count'], color='gray', alpha=0.7)
    ax.set_xlabel('Mutation Type')
    ax.set_ylabel('Count')
    ax.set_title('Sample Count by Mutation Type', fontweight='bold')
    ax.tick_params(axis='x', rotation=45)

    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

    return output_path, mutation_stats


def create_position_analysis(
    data: pd.DataFrame,
    output_path: str = None
) -> str:
    """
    Analyze delta by edit position.

    Args:
        data: DataFrame with edit_position and delta columns
        output_path: Path to save the plot

    Returns:
        Path to saved plot
    """
    set_plot_style()

    fig, axes = plt.subplots(1, 2, figsize=(12, 5))

    # Position distribution
    ax = axes[0]
    ax.hist(data['edit_position'], bins=50, color='steelblue', alpha=0.7, edgecolor='black')
    ax.set_xlabel('Edit Position (0-indexed)')
    ax.set_ylabel('Count')
    ax.set_title('Edit Position Distribution', fontweight='bold')

    # Delta vs Position (binned)
    ax = axes[1]

    # Bin positions
    n_bins = 10
    data = data.copy()
    data['pos_bin'] = pd.cut(data['edit_position'], bins=n_bins, labels=False)
    pos_stats = data.groupby('pos_bin')['delta'].agg(['mean', 'std']).reset_index()

    ax.errorbar(pos_stats['pos_bin'], pos_stats['mean'], yerr=pos_stats['std'],
                marker='o', linestyle='-', color='steelblue', capsize=5)
    ax.axhline(y=0, color='red', linestyle='--', alpha=0.7)
    ax.set_xlabel('Position Bin (5\' → 3\')')
    ax.set_ylabel('Mean Δ MRL')
    ax.set_title('Effect by Position', fontweight='bold')

    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

    return output_path


def create_edit_embedding_clustering(
    test_data: pd.DataFrame,
    embeddings: np.ndarray,
    output_path: str,
    n_clusters: int = 6
) -> Tuple[str, pd.DataFrame]:
    """
    Create edit embedding clustering visualization.

    Similar to small molecules, cluster the edit embeddings and analyze
    patterns in mutation effects.

    Args:
        test_data: DataFrame with edit information
        embeddings: Edit embeddings [n_samples, embedding_dim]
        output_path: Path to save the plot
        n_clusters: Number of clusters for K-means

    Returns:
        Tuple of (plot_path, cluster_stats_df)
    """
    from sklearn.cluster import KMeans
    from sklearn.decomposition import PCA

    set_plot_style()

    # Reduce dimensionality for visualization
    pca = PCA(n_components=2)
    emb_2d = pca.fit_transform(embeddings)

    # Cluster the embeddings
    kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
    cluster_labels = kmeans.fit_predict(embeddings)

    # Create figure with multiple subplots
    fig, axes = plt.subplots(2, 2, figsize=(14, 12))

    # 1. PCA scatter plot colored by cluster
    ax = axes[0, 0]
    scatter = ax.scatter(emb_2d[:, 0], emb_2d[:, 1], c=cluster_labels,
                        cmap='tab10', alpha=0.6, s=30)
    ax.set_xlabel('PC1')
    ax.set_ylabel('PC2')
    ax.set_title('Edit Embeddings (PCA) - Colored by Cluster', fontweight='bold')
    plt.colorbar(scatter, ax=ax, label='Cluster')

    # 2. PCA scatter plot colored by delta
    ax = axes[0, 1]
    scatter = ax.scatter(emb_2d[:, 0], emb_2d[:, 1], c=test_data['delta'],
                        cmap='RdYlGn_r', alpha=0.6, s=30, vmin=-1, vmax=1)
    ax.set_xlabel('PC1')
    ax.set_ylabel('PC2')
    ax.set_title('Edit Embeddings (PCA) - Colored by Δ MRL', fontweight='bold')
    plt.colorbar(scatter, ax=ax, label='Δ MRL')

    # 3. Delta distribution per cluster
    ax = axes[1, 0]
    test_data = test_data.copy()
    test_data['cluster'] = cluster_labels

    cluster_stats = []
    for cluster_id in range(n_clusters):
        cluster_data = test_data[test_data['cluster'] == cluster_id]
        cluster_stats.append({
            'cluster': cluster_id,
            'n_samples': len(cluster_data),
            'mean_delta': cluster_data['delta'].mean(),
            'std_delta': cluster_data['delta'].std(),
            'positive_frac': (cluster_data['delta'] > 0).mean()
        })

    stats_df = pd.DataFrame(cluster_stats)

    colors = plt.cm.tab10(np.linspace(0, 1, n_clusters))
    bars = ax.bar(stats_df['cluster'], stats_df['mean_delta'], color=colors,
                 alpha=0.8, edgecolor='black')
    ax.axhline(y=0, color='black', linestyle='-', linewidth=1)
    ax.set_xlabel('Cluster')
    ax.set_ylabel('Mean Δ MRL')
    ax.set_title('Mean Effect by Cluster', fontweight='bold')

    # Add error bars
    ax.errorbar(stats_df['cluster'], stats_df['mean_delta'],
               yerr=stats_df['std_delta'], fmt='none', color='black', capsize=5)

    # 4. Mutation type distribution per cluster
    ax = axes[1, 1]
    if 'edit_from' in test_data.columns and 'edit_to' in test_data.columns:
        test_data['mutation_type'] = test_data['edit_from'] + '→' + test_data['edit_to']

        # Get top mutation types
        top_mutations = test_data['mutation_type'].value_counts().head(6).index

        # Create stacked bar chart
        cluster_mutation_counts = []
        for cluster_id in range(n_clusters):
            cluster_data = test_data[test_data['cluster'] == cluster_id]
            counts = {'cluster': cluster_id}
            for mut in top_mutations:
                counts[mut] = (cluster_data['mutation_type'] == mut).sum()
            cluster_mutation_counts.append(counts)

        mutation_df = pd.DataFrame(cluster_mutation_counts)
        mutation_df.set_index('cluster', inplace=True)

        # Normalize to fractions
        mutation_df = mutation_df.div(mutation_df.sum(axis=1), axis=0)

        mutation_df.plot(kind='bar', stacked=True, ax=ax, colormap='Set2', alpha=0.8)
        ax.set_xlabel('Cluster')
        ax.set_ylabel('Fraction')
        ax.set_title('Mutation Type Distribution by Cluster', fontweight='bold')
        ax.legend(title='Mutation', bbox_to_anchor=(1.02, 1), loc='upper left', fontsize=8)
        ax.set_xticklabels(ax.get_xticklabels(), rotation=0)
    else:
        ax.text(0.5, 0.5, 'No mutation type data available',
               transform=ax.transAxes, ha='center', va='center', fontsize=12)
        ax.set_title('Mutation Type Distribution by Cluster', fontweight='bold')

    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

    return output_path, stats_df


def create_baseline_comparison_plot(
    results: Dict[str, Dict],
    data: pd.DataFrame,
    output_path: str
) -> str:
    """
    Compare methods against simple baselines.

    Args:
        results: Results dict with metrics
        data: DataFrame with delta values
        output_path: Path to save the plot

    Returns:
        Path to saved plot
    """
    set_plot_style()

    # Calculate baseline metrics
    delta_mean = data['delta'].mean()
    delta_std = data['delta'].std()

    # Mean predictor baseline (predict mean for all)
    mae_mean_baseline = np.mean(np.abs(data['delta'] - delta_mean))
    rmse_mean_baseline = np.sqrt(np.mean((data['delta'] - delta_mean) ** 2))

    # Zero predictor baseline (predict 0 for all)
    mae_zero_baseline = np.mean(np.abs(data['delta']))
    rmse_zero_baseline = np.sqrt(np.mean(data['delta'] ** 2))

    # Random baseline for direction accuracy
    direction_random = 0.5

    fig, axes = plt.subplots(1, 3, figsize=(15, 5))

    methods = list(results.keys())
    colors = plt.cm.Set2(np.linspace(0, 1, len(methods) + 2))

    # MAE comparison
    ax = axes[0]
    all_methods = methods + ['Mean Baseline', 'Zero Baseline']
    mae_values = [results[m].get('mae', 0) for m in methods] + [mae_mean_baseline, mae_zero_baseline]
    bars = ax.bar(all_methods, mae_values, color=colors)
    ax.set_ylabel('MAE')
    ax.set_title('MAE Comparison (Lower is Better)', fontweight='bold')
    ax.tick_params(axis='x', rotation=45)
    for bar, val in zip(bars, mae_values):
        ax.annotate(f'{val:.3f}', xy=(bar.get_x() + bar.get_width()/2, bar.get_height()),
                   xytext=(0, 3), textcoords="offset points", ha='center', fontsize=9)

    # RMSE comparison
    ax = axes[1]
    rmse_values = [results[m].get('rmse', 0) for m in methods] + [rmse_mean_baseline, rmse_zero_baseline]
    bars = ax.bar(all_methods, rmse_values, color=colors)
    ax.set_ylabel('RMSE')
    ax.set_title('RMSE Comparison (Lower is Better)', fontweight='bold')
    ax.tick_params(axis='x', rotation=45)
    for bar, val in zip(bars, rmse_values):
        ax.annotate(f'{val:.3f}', xy=(bar.get_x() + bar.get_width()/2, bar.get_height()),
                   xytext=(0, 3), textcoords="offset points", ha='center', fontsize=9)

    # Direction Accuracy comparison
    ax = axes[2]
    dir_methods = methods + ['Random Baseline']
    dir_values = [results[m].get('direction_accuracy', 0) for m in methods] + [direction_random]
    colors_dir = plt.cm.Set2(np.linspace(0, 1, len(dir_methods)))
    bars = ax.bar(dir_methods, dir_values, color=colors_dir)
    ax.axhline(y=0.5, color='red', linestyle='--', alpha=0.7, label='Random (50%)')
    ax.set_ylabel('Direction Accuracy')
    ax.set_title('Direction Accuracy (Higher is Better)', fontweight='bold')
    ax.set_ylim(0, 1)
    ax.tick_params(axis='x', rotation=45)
    for bar, val in zip(bars, dir_values):
        ax.annotate(f'{val:.2%}', xy=(bar.get_x() + bar.get_width()/2, bar.get_height()),
                   xytext=(0, 3), textcoords="offset points", ha='center', fontsize=9)

    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

    return output_path


def create_summary_table(
    results: Dict[str, Dict],
    data: pd.DataFrame
) -> pd.DataFrame:
    """
    Create summary table comparing all methods.

    Args:
        results: Results dict with metrics
        data: DataFrame for baseline calculations

    Returns:
        Summary DataFrame
    """
    # Calculate baselines
    delta_mean = data['delta'].mean()
    mae_mean = np.mean(np.abs(data['delta'] - delta_mean))
    rmse_mean = np.sqrt(np.mean((data['delta'] - delta_mean) ** 2))
    mae_zero = np.mean(np.abs(data['delta']))

    rows = []

    for method, metrics in results.items():
        rows.append({
            'Method': method,
            'MAE': f"{metrics.get('mae', 0):.4f}",
            'RMSE': f"{metrics.get('rmse', 0):.4f}",
            'R²': f"{metrics.get('r2', 0):.4f}",
            'Pearson r': f"{metrics.get('pearson_r', 0):.4f}",
            'Spearman ρ': f"{metrics.get('spearman_r', 0):.4f}",
            'Direction Acc': f"{metrics.get('direction_accuracy', 0):.2%}",
            'Train Size': metrics.get('train_size', 'N/A'),
            'Test Size': metrics.get('test_size', 'N/A')
        })

    # Add baselines
    rows.append({
        'Method': 'Mean Baseline',
        'MAE': f"{mae_mean:.4f}",
        'RMSE': f"{rmse_mean:.4f}",
        'R²': '0.0000',
        'Pearson r': 'N/A',
        'Spearman ρ': 'N/A',
        'Direction Acc': '50.00%',
        'Train Size': 'N/A',
        'Test Size': len(data)
    })

    rows.append({
        'Method': 'Zero Baseline',
        'MAE': f"{mae_zero:.4f}",
        'RMSE': f"{np.sqrt(np.mean(data['delta']**2)):.4f}",
        'R²': 'N/A',
        'Pearson r': 'N/A',
        'Spearman ρ': 'N/A',
        'Direction Acc': '50.00%',
        'Train Size': 'N/A',
        'Test Size': len(data)
    })

    return pd.DataFrame(rows)


def generate_rna_report(
    results: Dict[str, Dict],
    config: Any,
    data: Optional[pd.DataFrame] = None,
    predictions: Optional[Dict[str, Tuple[np.ndarray, np.ndarray]]] = None,
    trained_models: Optional[Dict] = None
) -> str:
    """
    Generate comprehensive RNA experiment report.

    Args:
        results: Dict of {method_name: {metric: value, ...}}
        config: RNAExperimentConfig object
        data: Optional DataFrame with test data
        predictions: Optional dict of {method_name: (y_true, y_pred)}
        trained_models: Optional dict of trained models (for training history)

    Returns:
        Path to generated report
    """
    output_dir = Path(config.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Create directories
    reports_dir = output_dir / "reports"
    images_dir = output_dir / "images" / timestamp
    reports_dir.mkdir(parents=True, exist_ok=True)
    images_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n{'='*60}")
    print("GENERATING RNA EXPERIMENT REPORT")
    print(f"{'='*60}")

    # Load data if not provided
    if data is None:
        data_path = Path(config.output_dir).parent.parent.parent / config.data_file
        if data_path.exists():
            data = pd.read_csv(data_path)
            print(f"Loaded data from: {data_path}")
        else:
            print(f"Warning: Could not load data from {data_path}")

    # 1. Create metrics comparison plot
    print("Creating metrics comparison plot...")
    metrics_to_plot = ['mae', 'rmse', 'r2', 'pearson_r', 'direction_accuracy']
    metrics_plot_path = create_metrics_comparison_plot(
        results, metrics_to_plot,
        str(images_dir / "metrics_comparison.png"),
        title=f"Method Comparison - {config.experiment_name}"
    )

    # 2. Create scatter plots (if predictions available)
    scatter_paths = {}
    if predictions:
        print("Creating scatter plots...")
        for method_name, (y_true, y_pred) in predictions.items():
            safe_name = method_name.replace(' ', '_').replace('-', '_')
            scatter_path = create_scatter_plot(
                y_true, y_pred, method_name,
                str(images_dir / f"scatter_{safe_name}.png")
            )
            scatter_paths[method_name] = scatter_path

    # 3. Create baseline comparison
    print("Creating baseline comparison...")
    if data is not None:
        baseline_path = create_baseline_comparison_plot(
            results, data,
            str(images_dir / "baseline_comparison.png")
        )

    # 4. Create delta distribution plot
    if data is not None:
        print("Creating delta distribution plot...")
        pred_deltas = {m: p[1] for m, p in predictions.items()} if predictions else None
        delta_dist_path = create_delta_distribution_plot(
            data, pred_deltas,
            str(images_dir / "delta_distribution.png")
        )

    # 5. Create mutation type analysis
    if data is not None and 'edit_from' in data.columns:
        print("Creating mutation type analysis...")
        mutation_path, mutation_stats = create_mutation_type_analysis(
            data, results,
            str(images_dir / "mutation_type_analysis.png")
        )

    # 6. Create position analysis
    if data is not None and 'edit_position' in data.columns:
        print("Creating position analysis...")
        position_path = create_position_analysis(
            data,
            str(images_dir / "position_analysis.png")
        )

    # 7. Create edit embedding clustering (if we have trained models)
    cluster_stats = None
    if trained_models is not None and data is not None:
        print("Creating edit embedding clustering...")
        try:
            # Try to extract edit embeddings from structured model
            for method_name, model_info in trained_models.items():
                if model_info.get('type') == 'structured_edit_framework':
                    model = model_info['model']
                    model.eval()

                    import torch
                    embeddings_list = []

                    # Get embeddings for test data
                    with torch.no_grad():
                        for _, row in data.iterrows():
                            emb = model.structured_embedder(
                                row['seq_a'],
                                row['edit_position'],
                                row['edit_from'],
                                row['edit_to']
                            )
                            if hasattr(emb, 'cpu'):
                                emb = emb.cpu().numpy()
                            embeddings_list.append(emb.flatten())

                    if embeddings_list:
                        embeddings = np.array(embeddings_list)
                        clustering_path, cluster_stats = create_edit_embedding_clustering(
                            data, embeddings,
                            str(images_dir / "edit_embedding_clustering.png"),
                            n_clusters=min(6, len(data) // 50)  # Adaptive clusters
                        )
                        print(f"  Created edit embedding clustering: {clustering_path}")
                    break
        except Exception as e:
            print(f"  Warning: Could not create edit embedding clustering: {e}")

    # 8. Create summary table
    if data is not None:
        summary_df = create_summary_table(results, data)
    else:
        summary_df = pd.DataFrame([
            {'Method': m, **{k: v for k, v in metrics.items()}}
            for m, metrics in results.items()
        ])

    # Save results as JSON
    results_file = output_dir / 'results.json'
    with open(results_file, 'w') as f:
        # Convert numpy types
        def convert(obj):
            if hasattr(obj, 'item'):
                return obj.item()
            return obj

        serializable = {
            k: {k2: convert(v2) for k2, v2 in v.items()}
            for k, v in results.items()
        }
        json.dump(serializable, f, indent=2)

    # Generate Markdown summary
    summary_path = output_dir / 'summary.md'
    _generate_markdown_summary(summary_path, results, config, summary_df, data)

    # Generate Word document (if docx available)
    if HAS_DOCX:
        docx_path = _generate_word_report(
            reports_dir, config, results, summary_df,
            images_dir, data, predictions, trained_models
        )
        print(f"\nWord report: {docx_path}")

    print(f"\n{'='*60}")
    print("REPORT GENERATION COMPLETE")
    print(f"{'='*60}")
    print(f"Images saved to: {images_dir}")
    print(f"Summary saved to: {summary_path}")
    print(f"Results saved to: {results_file}")

    return str(summary_path)


def _generate_markdown_summary(
    path: Path,
    results: Dict,
    config: Any,
    summary_df: pd.DataFrame,
    data: Optional[pd.DataFrame]
) -> None:
    """Generate markdown summary file."""
    with open(path, 'w') as f:
        f.write(f"# RNA Experiment Report: {config.experiment_name}\n\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")

        # Configuration
        f.write("## Configuration\n\n")
        f.write(f"- **Embedder**: {config.embedder_type}\n")
        f.write(f"- **Splitter**: {config.splitter_type}\n")
        f.write(f"- **Data file**: {config.data_file}\n")
        f.write(f"- **Train/Val/Test ratio**: {config.train_ratio}/{config.val_ratio}/{config.test_ratio}\n")
        f.write(f"- **Random seed**: {config.random_seed}\n\n")

        # Data Statistics
        if data is not None:
            f.write("## Data Statistics\n\n")
            f.write(f"- **Total pairs**: {len(data):,}\n")
            f.write(f"- **Unique sequences (seq_a)**: {data['seq_a'].nunique():,}\n")
            f.write(f"- **Delta mean**: {data['delta'].mean():.4f}\n")
            f.write(f"- **Delta std**: {data['delta'].std():.4f}\n")
            f.write(f"- **Delta range**: [{data['delta'].min():.4f}, {data['delta'].max():.4f}]\n\n")

        # Methods
        f.write("## Methods\n\n")
        for method in config.methods:
            f.write(f"### {method['name']}\n")
            f.write(f"- Type: {method['type']}\n")
            f.write(f"- Hidden dims: {method.get('hidden_dims', 'N/A')}\n")
            f.write(f"- Learning rate: {method.get('lr', 'N/A')}\n")
            f.write(f"- Batch size: {method.get('batch_size', 'N/A')}\n")
            f.write(f"- Max epochs: {method.get('max_epochs', 'N/A')}\n\n")

        # Results Table
        f.write("## Results\n\n")
        f.write(summary_df.to_markdown(index=False))
        f.write("\n\n")

        # Key Observations
        f.write("## Key Observations\n\n")

        # Find best method
        method_maes = {m: results[m].get('mae', float('inf')) for m in results}
        best_method = min(method_maes, key=method_maes.get)

        f.write(f"- **Best MAE**: {best_method} ({method_maes[best_method]:.4f})\n")

        # Direction accuracy analysis
        dir_accs = {m: results[m].get('direction_accuracy', 0) for m in results}
        best_dir = max(dir_accs, key=dir_accs.get)
        f.write(f"- **Best Direction Accuracy**: {best_dir} ({dir_accs[best_dir]:.2%})\n")

        # R² analysis
        r2s = {m: results[m].get('r2', float('-inf')) for m in results}
        best_r2 = max(r2s, key=r2s.get)
        f.write(f"- **Best R²**: {best_r2} ({r2s[best_r2]:.4f})\n\n")

        # Warnings
        f.write("## Potential Issues\n\n")
        if all(r2s[m] < 0.1 for m in r2s):
            f.write("- **Low R²**: All methods have R² < 0.1, indicating poor variance explanation\n")
        if all(dir_accs[m] < 0.6 for m in dir_accs):
            f.write("- **Low Direction Accuracy**: All methods near random chance (50%)\n")
        if data is not None and len(data) < 3000:
            f.write(f"- **Small Dataset**: Only {len(data):,} pairs - consider data augmentation\n")


def _generate_word_report(
    reports_dir: Path,
    config: Any,
    results: Dict,
    summary_df: pd.DataFrame,
    images_dir: Path,
    data: Optional[pd.DataFrame],
    predictions: Optional[Dict],
    trained_models: Optional[Dict]
) -> str:
    """Generate Word document report."""
    doc = Document()

    # Title
    title = doc.add_heading(f'RNA Experiment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Experiment: {config.experiment_name}')
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    # Configuration Section
    doc.add_heading('1. Experiment Configuration', 1)

    doc.add_heading('General Settings', 2)
    doc.add_paragraph(f"Embedder Type: {config.embedder_type}")
    doc.add_paragraph(f"Splitter Type: {config.splitter_type}")
    doc.add_paragraph(f"Data File: {config.data_file}")
    doc.add_paragraph(f"Train/Val/Test: {config.train_ratio}/{config.val_ratio}/{config.test_ratio}")
    doc.add_paragraph(f"Random Seed: {config.random_seed}")

    # Data Statistics
    if data is not None:
        doc.add_heading('Data Statistics', 2)
        doc.add_paragraph(f"Total pairs: {len(data):,}")
        doc.add_paragraph(f"Unique sequences: {data['seq_a'].nunique():,}")
        doc.add_paragraph(f"Delta mean ± std: {data['delta'].mean():.4f} ± {data['delta'].std():.4f}")
        doc.add_paragraph(f"Delta range: [{data['delta'].min():.4f}, {data['delta'].max():.4f}]")

    # Methods Configuration
    doc.add_heading('Methods', 2)
    for method in config.methods:
        doc.add_paragraph(f"• {method['name']}", style='List Bullet')
        doc.add_paragraph(f"  Type: {method['type']}")
        doc.add_paragraph(f"  Hidden dims: {method.get('hidden_dims', 'N/A')}")
        doc.add_paragraph(f"  LR: {method.get('lr', 'N/A')}, Batch: {method.get('batch_size', 'N/A')}")

    # Results Section
    doc.add_heading('2. Results', 1)

    # Add metrics comparison plot
    metrics_plot = images_dir / "metrics_comparison.png"
    if metrics_plot.exists():
        doc.add_heading('Metrics Comparison', 2)
        doc.add_picture(str(metrics_plot), width=Inches(6))

    # Add summary table
    doc.add_heading('Summary Table', 2)
    table = doc.add_table(rows=len(summary_df) + 1, cols=len(summary_df.columns))
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, col in enumerate(summary_df.columns):
        table.rows[0].cells[i].text = col
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    # Data rows
    for i, row in summary_df.iterrows():
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = str(val)

    # Baseline comparison
    baseline_plot = images_dir / "baseline_comparison.png"
    if baseline_plot.exists():
        doc.add_heading('Baseline Comparison', 2)
        doc.add_picture(str(baseline_plot), width=Inches(6))

    # Scatter plots
    if predictions:
        doc.add_heading('Prediction Scatter Plots', 2)
        for method_name in predictions.keys():
            safe_name = method_name.replace(' ', '_').replace('-', '_')
            scatter_path = images_dir / f"scatter_{safe_name}.png"
            if scatter_path.exists():
                doc.add_paragraph(f"{method_name}:")
                doc.add_picture(str(scatter_path), width=Inches(4))

    # Analysis Section
    doc.add_heading('3. Analysis', 1)

    # Delta distribution
    delta_dist = images_dir / "delta_distribution.png"
    if delta_dist.exists():
        doc.add_heading('Delta Distribution', 2)
        doc.add_picture(str(delta_dist), width=Inches(6))

    # Mutation type analysis
    mutation_plot = images_dir / "mutation_type_analysis.png"
    if mutation_plot.exists():
        doc.add_heading('Mutation Type Analysis', 2)
        doc.add_picture(str(mutation_plot), width=Inches(6))

    # Position analysis
    position_plot = images_dir / "position_analysis.png"
    if position_plot.exists():
        doc.add_heading('Position Analysis', 2)
        doc.add_picture(str(position_plot), width=Inches(6))

    # Edit Embedding Clustering
    clustering_plot = images_dir / "edit_embedding_clustering.png"
    if clustering_plot.exists():
        doc.add_heading('Edit Embedding Clustering', 2)
        doc.add_paragraph(
            "Clustering of learned edit embeddings reveals patterns in how mutations "
            "are represented and their relationship to expression effects."
        )
        doc.add_picture(str(clustering_plot), width=Inches(6))

    # Key Findings
    doc.add_heading('4. Key Findings', 1)

    method_maes = {m: results[m].get('mae', float('inf')) for m in results}
    best_method = min(method_maes, key=method_maes.get)
    doc.add_paragraph(f"• Best MAE: {best_method} ({method_maes[best_method]:.4f})", style='List Bullet')

    dir_accs = {m: results[m].get('direction_accuracy', 0) for m in results}
    best_dir = max(dir_accs, key=dir_accs.get)
    doc.add_paragraph(f"• Best Direction Accuracy: {best_dir} ({dir_accs[best_dir]:.2%})", style='List Bullet')

    r2s = {m: results[m].get('r2', float('-inf')) for m in results}
    best_r2 = max(r2s, key=r2s.get)
    doc.add_paragraph(f"• Best R²: {best_r2} ({r2s[best_r2]:.4f})", style='List Bullet')

    # Potential Issues
    doc.add_heading('Potential Issues', 2)
    if all(r2s[m] < 0.1 for m in r2s):
        doc.add_paragraph("• Low R²: All methods have R² < 0.1", style='List Bullet')
    if all(dir_accs[m] < 0.6 for m in dir_accs):
        doc.add_paragraph("• Low Direction Accuracy: Near random chance", style='List Bullet')
    if data is not None and len(data) < 3000:
        doc.add_paragraph(f"• Small Dataset: Only {len(data):,} pairs", style='List Bullet')

    # Save document
    report_path = reports_dir / f"{config.experiment_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(report_path))

    return str(report_path)


if __name__ == "__main__":
    # Test the report generator with existing results
    import argparse

    parser = argparse.ArgumentParser(description='Generate RNA experiment report')
    parser.add_argument('--results-dir', type=str, default='experiments/rna/results/random',
                       help='Path to results directory')
    args = parser.parse_args()

    results_dir = Path(args.results_dir)
    results_file = results_dir / 'results.json'

    if results_file.exists():
        with open(results_file) as f:
            results = json.load(f)

        # Create a minimal config for testing
        class MinimalConfig:
            experiment_name = "rna_mpra_test"
            embedder_type = "rnafm"
            splitter_type = "random"
            data_file = "../../data/rna/pairs/mpra_5utr_pairs_long.csv"
            train_ratio = 0.7
            val_ratio = 0.15
            test_ratio = 0.15
            random_seed = 42
            output_dir = str(results_dir)
            methods = [
                {'name': 'Structured Edit Framework', 'type': 'structured_edit_framework'},
                {'name': 'Edit Framework - RNA-FM', 'type': 'edit_framework'},
                {'name': 'Baseline Property Predictor', 'type': 'baseline_property'}
            ]

        config = MinimalConfig()

        # Load data
        data_path = Path("data/rna/pairs/mpra_5utr_pairs_long.csv")
        data = pd.read_csv(data_path) if data_path.exists() else None

        generate_rna_report(results, config, data)
    else:
        print(f"No results file found at {results_file}")
