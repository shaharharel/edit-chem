"""
Report generator for antibody mutation effect prediction experiments.

Generates both HTML and DOCX reports with:
- Experiment configuration summary
- Method comparison metrics
- Training progress visualization
- Prediction scatter plots
- Edit embedding analysis (t-SNE/UMAP)
- Cluster analysis
"""

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.manifold import TSNE
from sklearn.cluster import KMeans
from scipy import stats

# Optional DOCX support
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("python-docx not installed. DOCX reports will be skipped.")


class AntibodyReportGenerator:
    """Generate comprehensive reports for antibody experiments."""

    def __init__(
        self,
        output_dir: str,
        experiment_name: str,
        timestamp: Optional[str] = None,
    ):
        self.output_dir = Path(output_dir)
        self.experiment_name = experiment_name
        self.timestamp = timestamp or datetime.now().strftime("%Y%m%d_%H%M%S")

        # Create output directories
        self.reports_dir = self.output_dir / "reports"
        self.images_dir = self.output_dir / "images" / self.timestamp
        self.reports_dir.mkdir(parents=True, exist_ok=True)
        self.images_dir.mkdir(parents=True, exist_ok=True)

        # Collected data
        self.config = None
        self.results = {}
        self.embeddings = {}
        self.predictions = {}
        self.training_history = {}

    def set_config(self, config: Dict[str, Any]):
        """Set experiment configuration."""
        self.config = config

    def add_method_results(
        self,
        method_name: str,
        metrics: Dict[str, float],
        predictions: Optional[Tuple[np.ndarray, np.ndarray]] = None,
        embeddings: Optional[np.ndarray] = None,
        training_history: Optional[Dict[str, List[float]]] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ):
        """Add results for a method."""
        self.results[method_name] = {
            'metrics': metrics,
            'metadata': metadata or {},
        }

        if predictions is not None:
            y_true, y_pred = predictions
            self.predictions[method_name] = {
                'y_true': np.array(y_true),
                'y_pred': np.array(y_pred),
            }

        if embeddings is not None:
            self.embeddings[method_name] = np.array(embeddings)

        if training_history is not None:
            self.training_history[method_name] = training_history

    def generate_reports(
        self,
        generate_html: bool = True,
        generate_docx: bool = True,
    ) -> Dict[str, str]:
        """Generate all reports."""
        paths = {}

        # Generate plots first
        self._generate_all_plots()

        if generate_html:
            html_path = self._generate_html_report()
            paths['html'] = str(html_path)

        if generate_docx and HAS_DOCX:
            docx_path = self._generate_docx_report()
            paths['docx'] = str(docx_path)

        # Save results JSON
        json_path = self.reports_dir / f"{self.experiment_name}_{self.timestamp}_results.json"
        self._save_results_json(json_path)
        paths['json'] = str(json_path)

        return paths

    def _generate_all_plots(self):
        """Generate all visualization plots."""
        # Metrics comparison
        if self.results:
            self._plot_metrics_comparison()

        # Training curves
        for method_name, history in self.training_history.items():
            self._plot_training_curve(method_name, history)

        # Prediction scatter plots
        for method_name, preds in self.predictions.items():
            self._plot_predictions(method_name, preds['y_true'], preds['y_pred'])

        # Edit embedding analysis
        if self.embeddings:
            self._plot_embedding_analysis()

    def _plot_metrics_comparison(self):
        """Create metrics comparison bar chart."""
        methods = list(self.results.keys())
        metrics_names = ['mae', 'rmse', 'pearson', 'spearman', 'r2']

        fig, axes = plt.subplots(1, len(metrics_names), figsize=(4 * len(metrics_names), 4))
        if len(metrics_names) == 1:
            axes = [axes]

        colors = plt.cm.Set2(np.linspace(0, 1, len(methods)))

        for ax, metric in zip(axes, metrics_names):
            values = []
            for method in methods:
                val = self.results[method]['metrics'].get(metric, 0)
                values.append(val if val is not None else 0)

            bars = ax.bar(range(len(methods)), values, color=colors)
            ax.set_xticks(range(len(methods)))
            ax.set_xticklabels([m[:15] + '...' if len(m) > 15 else m for m in methods],
                               rotation=45, ha='right')
            ax.set_title(metric.upper())
            ax.grid(axis='y', alpha=0.3)

            # Add value labels
            for bar, val in zip(bars, values):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01,
                       f'{val:.3f}', ha='center', va='bottom', fontsize=8)

        plt.tight_layout()
        plt.savefig(self.images_dir / "metrics_comparison.png", dpi=150, bbox_inches='tight')
        plt.close()

    def _plot_training_curve(self, method_name: str, history: Dict[str, List[float]]):
        """Plot training curves."""
        fig, axes = plt.subplots(1, 2, figsize=(12, 4))

        # Loss curve
        if 'train_loss' in history:
            axes[0].plot(history['train_loss'], label='Train', color='blue')
        if 'val_loss' in history:
            axes[0].plot(history['val_loss'], label='Val', color='orange')
        axes[0].set_xlabel('Epoch')
        axes[0].set_ylabel('Loss')
        axes[0].set_title(f'{method_name} - Training Loss')
        axes[0].legend()
        axes[0].grid(alpha=0.3)

        # Metric curve (e.g., MAE or Pearson)
        if 'val_mae' in history:
            axes[1].plot(history['val_mae'], label='Val MAE', color='green')
        if 'val_pearson' in history:
            axes[1].plot(history.get('val_pearson', []), label='Val Pearson', color='purple')
        axes[1].set_xlabel('Epoch')
        axes[1].set_ylabel('Metric')
        axes[1].set_title(f'{method_name} - Validation Metrics')
        axes[1].legend()
        axes[1].grid(alpha=0.3)

        plt.tight_layout()
        safe_name = method_name.replace(' ', '_').replace('/', '_')
        plt.savefig(self.images_dir / f"training_{safe_name}.png", dpi=150, bbox_inches='tight')
        plt.close()

    def _plot_predictions(self, method_name: str, y_true: np.ndarray, y_pred: np.ndarray):
        """Plot prediction scatter plot."""
        fig, ax = plt.subplots(figsize=(8, 8))

        # Scatter plot
        ax.scatter(y_true, y_pred, alpha=0.5, s=20, edgecolors='none')

        # Perfect prediction line
        min_val = min(y_true.min(), y_pred.min())
        max_val = max(y_true.max(), y_pred.max())
        ax.plot([min_val, max_val], [min_val, max_val], 'r--', lw=2, label='Perfect')

        # Regression line
        slope, intercept = np.polyfit(y_true, y_pred, 1)
        ax.plot([min_val, max_val], [slope * min_val + intercept, slope * max_val + intercept],
                'g-', lw=1.5, label=f'Fit (slope={slope:.2f})')

        # Statistics
        r, p = stats.pearsonr(y_true, y_pred)
        mae = np.mean(np.abs(y_true - y_pred))
        rmse = np.sqrt(np.mean((y_true - y_pred) ** 2))

        ax.text(0.05, 0.95, f'Pearson r = {r:.3f}\nMAE = {mae:.3f}\nRMSE = {rmse:.3f}',
                transform=ax.transAxes, fontsize=10, verticalalignment='top',
                bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))

        ax.set_xlabel('True Values', fontsize=12)
        ax.set_ylabel('Predicted Values', fontsize=12)
        ax.set_title(f'{method_name}', fontsize=14)
        ax.legend(loc='lower right')
        ax.grid(alpha=0.3)

        plt.tight_layout()
        safe_name = method_name.replace(' ', '_').replace('/', '_')
        plt.savefig(self.images_dir / f"scatter_{safe_name}.png", dpi=150, bbox_inches='tight')
        plt.close()

    def _plot_embedding_analysis(self):
        """Analyze and visualize edit embeddings."""
        # Combine all embeddings
        all_embeddings = []
        labels = []
        for method_name, emb in self.embeddings.items():
            all_embeddings.append(emb)
            labels.extend([method_name] * len(emb))

        if not all_embeddings:
            return

        combined = np.vstack(all_embeddings)

        # t-SNE projection
        if len(combined) > 50:  # Only if enough samples
            tsne = TSNE(n_components=2, random_state=42, perplexity=min(30, len(combined) - 1))
            projected = tsne.fit_transform(combined)

            fig, ax = plt.subplots(figsize=(10, 10))

            unique_labels = list(set(labels))
            colors = plt.cm.Set1(np.linspace(0, 1, len(unique_labels)))

            for i, method in enumerate(unique_labels):
                mask = [l == method for l in labels]
                ax.scatter(projected[mask, 0], projected[mask, 1],
                          c=[colors[i]], label=method, alpha=0.6, s=30)

            ax.set_title('Edit Embedding t-SNE Projection', fontsize=14)
            ax.legend()
            ax.grid(alpha=0.3)

            plt.tight_layout()
            plt.savefig(self.images_dir / "embedding_tsne.png", dpi=150, bbox_inches='tight')
            plt.close()

        # Per-method embedding analysis with clustering
        for method_name, emb in self.embeddings.items():
            if len(emb) > 10:
                self._plot_embedding_clusters(method_name, emb)

    def _plot_embedding_clusters(self, method_name: str, embeddings: np.ndarray):
        """Cluster analysis of embeddings."""
        n_clusters = min(5, len(embeddings) // 10)
        if n_clusters < 2:
            return

        # Cluster
        kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
        clusters = kmeans.fit_predict(embeddings)

        # Project to 2D
        if embeddings.shape[1] > 2:
            tsne = TSNE(n_components=2, random_state=42, perplexity=min(30, len(embeddings) - 1))
            projected = tsne.fit_transform(embeddings)
        else:
            projected = embeddings

        fig, ax = plt.subplots(figsize=(8, 8))

        scatter = ax.scatter(projected[:, 0], projected[:, 1],
                            c=clusters, cmap='viridis', alpha=0.6, s=30)

        # Cluster centers (in projected space)
        for i in range(n_clusters):
            mask = clusters == i
            center = projected[mask].mean(axis=0)
            ax.scatter(center[0], center[1], c='red', s=200, marker='x',
                      linewidths=3, label=f'Cluster {i}' if i == 0 else '')

        ax.set_title(f'{method_name} - Edit Embedding Clusters', fontsize=14)
        ax.legend()
        ax.grid(alpha=0.3)

        plt.colorbar(scatter, ax=ax, label='Cluster')
        plt.tight_layout()

        safe_name = method_name.replace(' ', '_').replace('/', '_')
        plt.savefig(self.images_dir / f"clusters_{safe_name}.png", dpi=150, bbox_inches='tight')
        plt.close()

    def _generate_html_report(self) -> Path:
        """Generate HTML report."""
        html_path = self.reports_dir / f"{self.experiment_name}_{self.timestamp}.html"

        html_content = self._build_html()

        with open(html_path, 'w') as f:
            f.write(html_content)

        print(f"HTML report saved to: {html_path}")
        return html_path

    def _build_html(self) -> str:
        """Build HTML content."""
        # Convert results to JSON for JavaScript
        results_json = json.dumps({
            'config': self.config,
            'results': {
                method: {
                    'metrics': res['metrics'],
                    'metadata': res.get('metadata', {}),
                }
                for method, res in self.results.items()
            }
        }, default=str, indent=2)

        # Get relative image paths
        rel_images = self.images_dir.relative_to(self.reports_dir.parent)

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.experiment_name} - Antibody Experiment Report</title>
    <script src="https://cdn.plot.ly/plotly-2.27.0.min.js"></script>
    <style>
        :root {{
            --primary: #2563eb;
            --secondary: #64748b;
            --success: #10b981;
            --warning: #f59e0b;
            --danger: #ef4444;
            --bg-light: #f8fafc;
            --bg-card: #ffffff;
            --border: #e2e8f0;
            --text-primary: #1e293b;
            --text-secondary: #64748b;
        }}

        * {{ box-sizing: border-box; margin: 0; padding: 0; }}

        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: var(--bg-light);
            color: var(--text-primary);
            line-height: 1.6;
        }}

        .container {{ max-width: 1400px; margin: 0 auto; padding: 20px; }}

        header {{
            background: linear-gradient(135deg, var(--primary), #1d4ed8);
            color: white;
            padding: 40px 20px;
            margin-bottom: 30px;
            border-radius: 12px;
        }}

        header h1 {{ font-size: 2.5rem; margin-bottom: 10px; }}
        header p {{ opacity: 0.9; font-size: 1.1rem; }}

        .card {{
            background: var(--bg-card);
            border-radius: 12px;
            padding: 24px;
            margin-bottom: 24px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }}

        .card h2 {{
            font-size: 1.25rem;
            margin-bottom: 20px;
            padding-bottom: 12px;
            border-bottom: 1px solid var(--border);
        }}

        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 16px;
        }}

        .stat-item {{
            text-align: center;
            padding: 16px;
            background: var(--bg-light);
            border-radius: 8px;
        }}

        .stat-item .label {{
            font-size: 0.85rem;
            color: var(--text-secondary);
            text-transform: uppercase;
        }}

        .stat-item .value {{
            font-size: 1.5rem;
            font-weight: 600;
            color: var(--primary);
        }}

        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 16px 0;
        }}

        th, td {{
            padding: 12px;
            text-align: left;
            border-bottom: 1px solid var(--border);
        }}

        th {{
            background: var(--bg-light);
            font-weight: 600;
            color: var(--text-secondary);
        }}

        tr:hover {{ background: var(--bg-light); }}

        .plot-container {{
            width: 100%;
            margin: 20px 0;
        }}

        .plot-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(400px, 1fr));
            gap: 20px;
        }}

        .plot-item {{ text-align: center; }}
        .plot-item img {{ max-width: 100%; height: auto; border-radius: 8px; }}
        .plot-item .caption {{ margin-top: 8px; color: var(--text-secondary); font-size: 0.9rem; }}

        .best-metric {{ background-color: #dcfce7; font-weight: bold; }}

        footer {{
            text-align: center;
            padding: 30px;
            color: var(--text-secondary);
        }}
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>{self.experiment_name}</h1>
            <p>Antibody Mutation Effect Prediction Experiment Report</p>
            <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </header>

        <div class="card">
            <h2>Experiment Configuration</h2>
            <div class="stats-grid">
                {self._html_config_stats()}
            </div>
        </div>

        <div class="card">
            <h2>Method Comparison</h2>
            {self._html_metrics_table()}
        </div>

        <div class="card">
            <h2>Metrics Visualization</h2>
            <div class="plot-container">
                <img src="{rel_images}/metrics_comparison.png" alt="Metrics Comparison" style="max-width: 100%;">
            </div>
        </div>

        <div class="card">
            <h2>Prediction Scatter Plots</h2>
            <div class="plot-grid">
                {self._html_scatter_plots(rel_images)}
            </div>
        </div>

        <div class="card">
            <h2>Training Progress</h2>
            <div class="plot-grid">
                {self._html_training_plots(rel_images)}
            </div>
        </div>

        {self._html_embedding_section(rel_images)}

        <footer>
            <p>Generated by edit-chem antibody experiments pipeline</p>
        </footer>
    </div>

    <script>
        const experimentData = {results_json};
        console.log('Experiment data:', experimentData);
    </script>
</body>
</html>
"""
        return html

    def _html_config_stats(self) -> str:
        """Generate HTML for config stats."""
        if not self.config:
            return "<p>No configuration available</p>"

        stats = []
        if 'data_file' in self.config:
            stats.append(('Data File', Path(self.config['data_file']).name))
        if 'source_datasets' in self.config:
            stats.append(('Sources', ', '.join(self.config.get('source_datasets', ['all']))))
        if 'train_ratio' in self.config:
            stats.append(('Split', f"{self.config['train_ratio']}/{self.config.get('val_ratio', 0.15)}/{self.config.get('test_ratio', 0.15)}"))
        if 'splitter_type' in self.config:
            stats.append(('Splitter', self.config['splitter_type']))
        stats.append(('Methods', str(len(self.results))))

        html = ""
        for label, value in stats:
            html += f"""
            <div class="stat-item">
                <div class="label">{label}</div>
                <div class="value">{value}</div>
            </div>
            """
        return html

    def _html_metrics_table(self) -> str:
        """Generate HTML metrics comparison table."""
        if not self.results:
            return "<p>No results available</p>"

        # Get all metrics
        all_metrics = set()
        for res in self.results.values():
            all_metrics.update(res['metrics'].keys())
        metrics = sorted(all_metrics)

        # Find best values for highlighting
        best_values = {}
        for metric in metrics:
            values = [(method, res['metrics'].get(metric, None))
                     for method, res in self.results.items()
                     if res['metrics'].get(metric) is not None]
            if values:
                # Lower is better for mae, rmse; higher for others
                if metric in ['mae', 'rmse', 'mse']:
                    best_values[metric] = min(values, key=lambda x: x[1])[0]
                else:
                    best_values[metric] = max(values, key=lambda x: x[1])[0]

        # Build table
        html = "<table><thead><tr><th>Method</th>"
        for metric in metrics:
            html += f"<th>{metric.upper()}</th>"
        html += "</tr></thead><tbody>"

        for method, res in self.results.items():
            html += f"<tr><td><strong>{method}</strong></td>"
            for metric in metrics:
                val = res['metrics'].get(metric, '-')
                if isinstance(val, (int, float)):
                    cell_class = 'best-metric' if best_values.get(metric) == method else ''
                    html += f"<td class='{cell_class}'>{val:.4f}</td>"
                else:
                    html += f"<td>{val}</td>"
            html += "</tr>"

        html += "</tbody></table>"
        return html

    def _html_scatter_plots(self, rel_images: Path) -> str:
        """Generate HTML for scatter plots."""
        html = ""
        for method in self.predictions.keys():
            safe_name = method.replace(' ', '_').replace('/', '_')
            img_path = rel_images / f"scatter_{safe_name}.png"
            if (self.reports_dir.parent / img_path).exists():
                html += f"""
                <div class="plot-item">
                    <img src="{img_path}" alt="{method} predictions">
                    <div class="caption">{method}</div>
                </div>
                """
        return html or "<p>No prediction plots available</p>"

    def _html_training_plots(self, rel_images: Path) -> str:
        """Generate HTML for training plots."""
        html = ""
        for method in self.training_history.keys():
            safe_name = method.replace(' ', '_').replace('/', '_')
            img_path = rel_images / f"training_{safe_name}.png"
            if (self.reports_dir.parent / img_path).exists():
                html += f"""
                <div class="plot-item">
                    <img src="{img_path}" alt="{method} training">
                    <div class="caption">{method}</div>
                </div>
                """
        return html or "<p>No training plots available</p>"

    def _html_embedding_section(self, rel_images: Path) -> str:
        """Generate HTML for embedding analysis section."""
        if not self.embeddings:
            return ""

        html = """
        <div class="card">
            <h2>Edit Embedding Analysis</h2>
        """

        # t-SNE plot
        tsne_path = rel_images / "embedding_tsne.png"
        if (self.reports_dir.parent / tsne_path).exists():
            html += f"""
            <h3>t-SNE Projection</h3>
            <div class="plot-container">
                <img src="{tsne_path}" alt="Embedding t-SNE" style="max-width: 600px;">
            </div>
            """

        # Cluster plots
        html += "<h3>Cluster Analysis</h3><div class='plot-grid'>"
        for method in self.embeddings.keys():
            safe_name = method.replace(' ', '_').replace('/', '_')
            cluster_path = rel_images / f"clusters_{safe_name}.png"
            if (self.reports_dir.parent / cluster_path).exists():
                html += f"""
                <div class="plot-item">
                    <img src="{cluster_path}" alt="{method} clusters">
                    <div class="caption">{method}</div>
                </div>
                """
        html += "</div></div>"

        return html

    def _generate_docx_report(self) -> Path:
        """Generate DOCX report."""
        if not HAS_DOCX:
            return None

        docx_path = self.reports_dir / f"{self.experiment_name}_{self.timestamp}.docx"

        doc = Document()

        # Title
        title = doc.add_heading(self.experiment_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Antibody Mutation Effect Prediction Experiment Report")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Configuration
        doc.add_heading('Experiment Configuration', 1)
        if self.config:
            for key, value in self.config.items():
                if not isinstance(value, (dict, list)):
                    doc.add_paragraph(f"{key}: {value}")

        # Results table
        doc.add_heading('Method Comparison', 1)

        if self.results:
            # Get metrics
            all_metrics = set()
            for res in self.results.values():
                all_metrics.update(res['metrics'].keys())
            metrics = sorted(all_metrics)

            # Create table
            table = doc.add_table(rows=len(self.results) + 1, cols=len(metrics) + 1)
            table.style = 'Light Grid Accent 1'

            # Header
            table.rows[0].cells[0].text = 'Method'
            for i, metric in enumerate(metrics):
                table.rows[0].cells[i + 1].text = metric.upper()

            # Data
            for row_idx, (method, res) in enumerate(self.results.items(), 1):
                table.rows[row_idx].cells[0].text = method
                for col_idx, metric in enumerate(metrics):
                    val = res['metrics'].get(metric, '-')
                    if isinstance(val, (int, float)):
                        table.rows[row_idx].cells[col_idx + 1].text = f"{val:.4f}"
                    else:
                        table.rows[row_idx].cells[col_idx + 1].text = str(val)

        # Metrics comparison plot
        doc.add_heading('Metrics Visualization', 1)
        metrics_img = self.images_dir / "metrics_comparison.png"
        if metrics_img.exists():
            doc.add_picture(str(metrics_img), width=Inches(6))

        # Prediction plots
        doc.add_heading('Prediction Scatter Plots', 1)
        for method in self.predictions.keys():
            safe_name = method.replace(' ', '_').replace('/', '_')
            img_path = self.images_dir / f"scatter_{safe_name}.png"
            if img_path.exists():
                doc.add_heading(method, 2)
                doc.add_picture(str(img_path), width=Inches(4))

        # Training plots
        doc.add_heading('Training Progress', 1)
        for method in self.training_history.keys():
            safe_name = method.replace(' ', '_').replace('/', '_')
            img_path = self.images_dir / f"training_{safe_name}.png"
            if img_path.exists():
                doc.add_heading(method, 2)
                doc.add_picture(str(img_path), width=Inches(6))

        # Embedding analysis
        if self.embeddings:
            doc.add_heading('Edit Embedding Analysis', 1)

            tsne_path = self.images_dir / "embedding_tsne.png"
            if tsne_path.exists():
                doc.add_heading('t-SNE Projection', 2)
                doc.add_picture(str(tsne_path), width=Inches(5))

            for method in self.embeddings.keys():
                safe_name = method.replace(' ', '_').replace('/', '_')
                cluster_path = self.images_dir / f"clusters_{safe_name}.png"
                if cluster_path.exists():
                    doc.add_heading(f'{method} - Clusters', 2)
                    doc.add_picture(str(cluster_path), width=Inches(4))

        # Save
        doc.save(str(docx_path))
        print(f"DOCX report saved to: {docx_path}")

        return docx_path

    def _save_results_json(self, path: Path):
        """Save results to JSON file."""
        output = {
            'experiment_name': self.experiment_name,
            'timestamp': self.timestamp,
            'config': self.config,
            'results': {
                method: {
                    'metrics': res['metrics'],
                    'metadata': res.get('metadata', {}),
                }
                for method, res in self.results.items()
            },
        }

        with open(path, 'w') as f:
            json.dump(output, f, indent=2, default=str)

        print(f"Results JSON saved to: {path}")
